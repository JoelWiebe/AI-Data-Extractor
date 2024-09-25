import os
import json
import pandas as pd
import docx
from google.cloud import aiplatform
import datetime
from config import *
import re
import vertexai
from vertexai.generative_models import GenerativeModel




class ParagraphClassifierClient:
    def __init__(self):
        vertexai.init(project=PROJECT_ID, location=LOCATION)
        self.model = GenerativeModel(GEMINI_MODEL,
            system_instruction="""You are a meticulous research assistant with expertise in natural language processing. You will be assigned two main tasks:
            1. **Paragraph Classification:** Given a research paper section heading and its paragraphs, you will classify each paragraph based on a set of predefined labels, along with your confidence in each label. You will be provided with descriptions of these labels to guide your classification.
            2. **Variable Extraction:** Given a research paper section heading, paragraphs, and a list of target variables, you will extract the values of these variables from the paragraphs. For each extracted value, you will provide a justification explaining how you derived it from the text, referencing the most relevant paragraph(s). You will be provided with detailed descriptions of the target variables to help you accurately identify and extract them."""
            )

    def classify_section(self, heading, paragraphs, start_idx):
        print(f"\nProcessing section: {heading}")

        # Filter out blank or non-text paragraphs
        filtered_paragraphs = {str(i + start_idx): p 
                                for i, p in enumerate(paragraphs) 
                                if not p.isspace() and p.strip()}
        
        # Create JSON payload for Gemini with absolute paragraph indices
        payload = {
            "heading": heading,
            "paragraphs": filtered_paragraphs
        }

        json_payload = json.dumps(payload, indent=2)

        # Prompt template with instructions and expected response format (including confidence)
        prompt = (
            "The following is a JSON object containing a section heading and its paragraphs from a research paper. For the paragraphs dictionary, the keys are paragraph indices, and the values are the paragraph text:\n\n"
            f"{json_payload}\n\n"
            f"Next are the paragraph labels, along with their descriptions, that you will use to classify each of the paragraphs above:\n\n"
            f"{PARAGRAPH_TAG_DESCRIPTIONS}\n\n"
            "If a paragraph is relevant to multiple labels, assign multiple labels. "
            "For each relevant paragraph, provide a list of the assigned label names and your confidence (0 to 1) in each label. "
            "If a paragraph is not relevant to any label, do not include it in the response.\n\n"
            "The response should be a JSON object with the following format:\n\n"
            "{\n  \"classifications\": {\n    \"[Paragraph Index]\": [[\"label_name\", confidence], ...]\n  }\n}\n\n"
            "Do not include any Markdown formatting in your response. Ensure the JSON is valid and does not contain any extra characters or formatting.\n"
        )
        print(f"\nClassify section prompt:\n\n{prompt}")

        # Generate classifications 
        generation_config = GENERATION_CONFIGURATION

        safety_settings = SAFETY_SETTINGS
        
        # Call the model to predict and get results in string format
        response = self.model.generate_content([prompt], generation_config=generation_config, safety_settings=safety_settings).text
        print(f"\nClassify section response:\n\n{response}")

        clean_response = remove_json_markdown(response)

        # convert the results to a python dictionary
        response_dict = json.loads(clean_response)

        # updated parsing with confidence scores
        return response_dict["classifications"] # Return the dictionary
    

    def extract_target_variables(self, classified_paragraphs_data):
        extraction_results = {}
        print(f"Classified paragraph data:\n{classified_paragraphs_data}")

        for tag, heading_data in classified_paragraphs_data.items():
            if tag in TARGET_VARIABLES or tag in CLUSTER_TARGET_VARIABLES:
                target_variables_subset = {tag: TARGET_VARIABLES[tag]} if tag in TARGET_VARIABLES else {
                    var_name: TARGET_VARIABLES[var_name] for var_name in CLUSTER_TARGET_VARIABLES[tag]}

                # Prepare paragraphs under their respective headings
                paragraphs_by_heading = {}  
                for heading, paragraph_data in heading_data.items():
                    paragraphs_by_heading[heading] = {
                        str(idx): text for _, idx, text in paragraph_data
                    }

                # Create the JSON payload with paragraphs organized by heading
                payload = {
                    "target_variables": target_variables_subset,
                    "headings": paragraphs_by_heading
                }

                json_payload = json.dumps(payload, indent=2)

                # Craft the Gemini prompt
                prompt_start = f"You are given the following JSON object containing paragraphs from a research paper under different headings, and a list of target variables:\n\n{json_payload}\n\nPlease extract the values for the specified target variables from the paragraphs.\nIf a value cannot be found with high confidence, indicate 'Not Found'.\n\nFor each extracted value, provide a brief justification explaining how you deduced it from the text, referencing the most relevant paragraph(s).\n\nReturn your results in this JSON format:\n{{\n"

                prompt_end = "\n}\nDo not include any Markdown formatting in your response. Ensure the JSON is valid and does not contain any extra characters or formatting.\n"

                # Use a list comprehension to create the variable parts of the prompt, now including justification
                variable_parts = [
                    f'  "{var_name}": ' + '{' + f'"value": "[extracted value]", "confidence": [0-1], "indices": [list of relevant paragraph indices], "justification": "[brief explanation]"' + '}' 
                    for var_name in target_variables_subset
                ]

                # Join the variable parts with commas
                prompt_variables = ",\n".join(variable_parts)

                # Combine the prompt parts
                prompt = prompt_start + prompt_variables + prompt_end

                print(f"\nExtraction prompt:\n\n{prompt}")

                generation_config = GENERATION_CONFIGURATION

                safety_settings = SAFETY_SETTINGS

                # Generate and parse response (remember to remove any markdown)
                response = self.model.generate_content([prompt], generation_config=generation_config, safety_settings=safety_settings).text
                print(f"\nExtraction response:\n\n{response}")  # Print results for debugging
                clean_response = remove_json_markdown(response)
                response_dict = json.loads(clean_response)

                # Convert indices from string to integer
                for var_name, result in response_dict.items():
                    if 'indices' in result:
                        result['indices'] = [int(idx) for idx in result['indices']]
                
                # Update extraction results with individual target variables
                for var_name, result in response_dict.items():
                    extraction_results[var_name] = result

        print(f"Extraction results:\n{extraction_results}")
        return extraction_results

  
def remove_json_markdown(text):
    pattern = re.compile(r'```json\s*(.*?)\s*```', re.DOTALL)
    return pattern.sub(r'\1', text)


def update_classified_data(doc, par_classifier_client, classified_paragraphs_data, current_heading, current_paragraphs, start_idx):
    classifications = par_classifier_client.classify_section(current_heading, current_paragraphs, start_idx)  
                
    # Iterate through the classified paragraphs
    for temp_par_index, labels_with_confidences in classifications.items():
        # Add results to classified paragraphs data dictionary
        for label, confidence in labels_with_confidences:
            if label in PARAGRAPH_TAG_DESCRIPTIONS:
                # If the label is not in classified_paragraphs_data yet, initialize it
                if label not in classified_paragraphs_data:
                    classified_paragraphs_data[label] = {}
                # If the heading is not in the label's dictionary yet, initialize it
                if current_heading not in classified_paragraphs_data[label]:
                    classified_paragraphs_data[label][current_heading] = []
                # Add relevant paragraph classification data to the list under the corresponding heading
                classified_paragraphs_data[label][current_heading].append(
                    (confidence, int(temp_par_index), doc.paragraphs[int(temp_par_index)].text)
                )
            else:
                print(f"ERROR: label {label} is not a paragraph tag (will not be stored)")

    
def process_document(file_path, par_classifier_client):
    """
    Reads a Word document, classifies paragraphs into sections based on headings,
    and aggregates relevant paragraphs with their confidence scores for each target variable.

    Args:
        file_path (str): The path to the Word document.
        par_classifier_client (TagClassifierClient): The client for classifying paragraphs.

    Returns:
        dict: A dictionary where keys are target variable names and values are lists of tuples 
              containing paragraph indices and their corresponding confidence scores.
    """
    
    doc = docx.Document(file_path)

    # Initialize variables to track the current section and paragraph index
    current_heading = None  
    current_paragraphs = []
    paragraph_index = 0

    # Use tags as level 1 keys, then headings as level 2
    classified_paragraphs_data = {par_tag: {} for par_tag in PARAGRAPH_TAG_DESCRIPTIONS.keys()} 

    for para in doc.paragraphs:
        # Check if the current paragraph is a heading
        if para.style.name == 'Heading 1' or para.style.name == 'Heading 2':
            # If we have a previous section (heading and paragraphs), process it
            if current_heading and current_paragraphs:
                
                start_idx = paragraph_index - len(current_paragraphs)
                # Classify the paragraphs in the current section using the Gemini model
                update_classified_data(doc, par_classifier_client, classified_paragraphs_data, current_heading, current_paragraphs, start_idx)
                
                current_paragraphs = []

            # Update the current heading to the new heading
            current_heading = para.text
        else:
            # If the current paragraph is not a heading, add it to the current_paragraphs list
            current_paragraphs.append(para.text)
            start_idx = paragraph_index - len(current_paragraphs)
        # Increment the paragraph index for the next paragraph
        paragraph_index += 1

    # Process the last section of the document (if any)
    if current_heading and current_paragraphs:
        start_idx += 1
        update_classified_data(doc, par_classifier_client, classified_paragraphs_data, current_heading, current_paragraphs, start_idx) 

    # Return the dictionary containing aggregated data for all target variables
    return classified_paragraphs_data


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)  # Create output directory if not exists
    par_classifier_client = ParagraphClassifierClient()

    all_results = []  # List to store results from all documents

    for filename in os.listdir(INPUT_DIR):
        file_path = os.path.join(INPUT_DIR, filename)
        if os.path.isfile(file_path) and file_path.endswith(".docx") and not filename.startswith("~"):
            try:
                classified_paragraph_data = process_document(file_path, par_classifier_client)

                extracted_results = par_classifier_client.extract_target_variables(classified_paragraph_data)
                
                # Load the document to retrieve paragraph text
                doc = docx.Document(file_path) 

                # Convert extracted_results to the desired format
                for var_name, extraction_info in extracted_results.items():
                    relevant_paragraphs = []
                    for idx in extraction_info["indices"]:
                        relevant_paragraphs.append(f"{idx}: {doc.paragraphs[idx].text}")
                    
                    all_results.append({
                        "filename": filename,
                        "variable": var_name,
                        "relevant_paragraphs": "\n".join(relevant_paragraphs),
                        "extracted_value": extraction_info["value"],
                        "justification": extraction_info["justification"],
                        "human_verified_response": ""  # Empty cell for human verification
                    })

            except Exception as e:
                print(f"\nError processing {filename}: {e}\n")

    # Create DataFrame and save to Excel with timestamp
    df = pd.DataFrame(all_results)

    now = datetime.datetime.now()
    timestamp = now.strftime("%Y-%m-%d_%H-%M-%S")

    output_file = os.path.join(OUTPUT_DIR, f"extracted_data_{timestamp}.xlsx")
    df.to_excel(output_file, index=False)
    print(f"\nResults saved to: {output_file}")


if __name__ == "__main__":
    main()
