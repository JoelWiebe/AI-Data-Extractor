# ai-data-extractor.py

import os
import json
import pandas as pd
import docx
from docx.document import Document as DocxDocument 
from docx.table import Table as DocxTable 
import datetime
from config import * 
import re
import vertexai
from vertexai.generative_models import GenerativeModel
import time 
import sys
from google.api_core import exceptions as google_exceptions 


# Helper function to convert a python-docx Table object to GitHub Flavored Markdown
def docx_table_to_markdown(table_obj: DocxTable) -> str:
    """Converts a python-docx Table object to a GitHub Flavored Markdown string."""
    md_rows = []
    for i, row in enumerate(table_obj.rows):
        # Escape pipe characters within cell text to avoid breaking Markdown table structure
        cells_text = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
        md_rows.append("| " + " | ".join(cells_text) + " |")
        if i == 0:  # After header row, add separator
            separator = ["---"] * len(cells_text)
            md_rows.append("| " + " | ".join(separator) + " |")
    return "\n".join(md_rows)


class ParagraphClassifierClient:
    def __init__(self):
        vertexai.init(project=PROJECT_ID, location=LOCATION)
        self.model = GenerativeModel(GEMINI_MODEL,
            system_instruction="""You are a meticulous research assistant with expertise in natural language processing. Your primary focus will be on analyzing the methodologies, findings, and details of **the main, current research study being reported in the provided academic articles.** You will be assigned two main tasks:
            1. **Paragraph Classification:** Given a research paper section heading and its paragraphs (which may include text paragraphs or tables formatted as Markdown), you will classify each paragraph/table based on a set of predefined labels, along with your confidence in each label. You will be provided with descriptions of these labels to guide your classification.
            2. **Variable Extraction:** Given a research paper section heading, paragraphs (which may include text paragraphs or tables formatted as Markdown), and a list of target variables, you will extract the values of these variables from the paragraphs/tables. For each extracted value, you will provide a justification explaining how you derived it from the text, referencing the most relevant paragraph(s)/table(s). You will be provided with detailed descriptions of the target variables to help you accurately identify and extract them."""
            )
        
    def _handle_llm_response_issues(self, response_obj, task_description):
        """
        Checks for issues like MAX_TOKENS or SAFETY in the response candidate.
        If MAX_TOKENS, attempts to include any partially received text in the error.
        """
        if not response_obj.candidates:
            raise ValueError(f"{task_description} failed: Response object has no candidates.")
        
        candidate = response_obj.candidates[0]
        finish_reason_name = candidate.finish_reason.name if candidate.finish_reason else "UNKNOWN"

        # Attempt to extract any partial text, regardless of finish_reason initially
        # This helps in logging what was received, even if an error is raised later.
        partial_text_received = "[No text parts found in candidate content]" # Default message
        if candidate.content and candidate.content.parts:
            try:
                # Join text from all available parts
                collected_parts_text = "".join(part.text for part in candidate.content.parts if hasattr(part, 'text') and part.text is not None)
                if collected_parts_text:
                    partial_text_received = collected_parts_text
                # If collected_parts_text is empty, the default message remains.
            except Exception as e_text_extraction:
                partial_text_received = f"[Error attempting to extract partial text: {str(e_text_extraction)}]"
        
        # Now check for specific finish reasons
        if finish_reason_name == "MAX_TOKENS":
            error_message = (
                f"{task_description} stopped due to MAX_TOKENS. Output was truncated.\n"
                f"Partially Received Text: <<<{partial_text_received}>>>\n"
                f"Full Candidate Details (for debugging): {candidate}"
            )
            raise ValueError(error_message)

        if finish_reason_name == "SAFETY":
            error_message = (
                f"{task_description} blocked by safety filters.\n"
                f"Partially Received Text (if any, typically none for SAFETY block): <<<{partial_text_received}>>>\n"
                f"Safety Ratings: {candidate.safety_ratings}.\n"
                f"Full Candidate Details (for debugging): {candidate}"
            )
            raise ValueError(error_message)

        if finish_reason_name not in ["STOP", "FINISH_REASON_UNSPECIFIED"]:
            # For other unexpected finish reasons
            error_message = (
                f"{task_description} finished with unexpected reason: {finish_reason_name}.\n"
                f"Received Text (if any): <<<{partial_text_received}>>>\n"
                f"Full Candidate Details (for debugging): {candidate}"
            )
            raise ValueError(error_message)
        
        # If finish_reason is STOP or UNSPECIFIED, but the extracted text is effectively empty
        # (which implies candidate.content or candidate.content.parts was empty or had no text)
        if not partial_text_received or partial_text_received == "[No text parts found in candidate content]":
             # This check handles cases where finish_reason is STOP/UNSPECIFIED but content is still missing.
            error_message = (
                f"{task_description} response candidate content appears empty or has no text, "
                f"despite finish_reason being '{finish_reason_name}'.\n"
                f"Attempted Text Extraction: <<<{partial_text_received}>>>\n"
                f"Full Candidate Details (for debugging): {candidate}"
            )
            raise ValueError(error_message)

        # If all checks pass and text was successfully extracted (i.e., partial_text_received is the full text)
        return partial_text_received

    def classify_section(self, heading: str, section_content_strings: list[str], section_global_start_idx: int) -> dict:
        """
        Classifies a list of content strings (paragraphs or Markdown tables) under a heading.
        Args:
            heading (str): The heading of the section.
            section_content_strings (list[str]): List of text paragraphs or Markdown table strings.
            section_global_start_idx (int): The global index (in the document's full content list) 
                                             of the first string in section_content_strings.
        Returns:
            dict: Classifications keyed by global paragraph/content index (as strings), 
                  or an empty dict if no content or if a persistent error occurs after retries.
        Raises:
            RuntimeError: If classification fails after all retry attempts.
        """
        task_description = f"Classification for section '{heading}'"
        print(f"\n{task_description} (Content pieces: {len(section_content_strings)}, Global start idx: {section_global_start_idx})")

        # Create a dictionary for the payload where keys are global indices (as strings)
        # and values are the content strings (paragraph text or Markdown table).
        payload_paragraphs = {
            str(local_idx + section_global_start_idx): content_str
            for local_idx, content_str in enumerate(section_content_strings)
            if content_str and not content_str.isspace() # Ensure content is not just whitespace
        }

        if not payload_paragraphs:
            print(f"No non-empty content to classify in section: {heading}")
            return {}

        payload = {
            "heading": heading,
            "paragraphs": payload_paragraphs # Keys are global indices as strings
        }
        json_payload_for_prompt = json.dumps(payload, indent=2) # For inclusion in the prompt

        # Prepare prompt components
        valid_label_names = list(PARAGRAPH_TAG_DESCRIPTIONS.keys())
        
        formatted_descriptions = "\n\nAvailable Labels and What They Cover:\n"
        for label_name, description_list in PARAGRAPH_TAG_DESCRIPTIONS.items():
            joined_descriptions = "; ".join(description_list) 
            formatted_descriptions += f"- **{label_name}**: This label pertains to content about: {joined_descriptions}\n"

        # Construct the full prompt
        prompt = (
            "The following is a JSON object containing a section heading and its associated content pieces (paragraphs or tables formatted as Markdown) from a research paper. For the 'paragraphs' dictionary, the keys are unique content piece indices (as strings), and the values are the content string (text or Markdown table):\n\n"
            f"{json_payload_for_prompt}\n\n"
            "Your task is to classify each content piece. You MUST ONLY use label names from the following predefined list:\n"
            f"VALID LABEL NAMES: [{', '.join(valid_label_names)}]\n\n" 
            f"To help you understand what each valid label name means, refer to these descriptions:\n"
            f"{formatted_descriptions}\n\n"
            "If a content piece is relevant to multiple labels, assign multiple labels using ONLY names from the VALID LABEL NAMES list. "
            "For each relevant content piece, provide a list containing pairs of [\"exact_label_name_from_valid_list\", confidence_score_0_to_1]. "
            "If a content piece is not relevant to any of the listed VALID LABEL NAMES, do not include that content piece index in your response's 'classifications' object.\n\n"
            "The response MUST be a single JSON object with the following format. Pay EXTREMELY close attention to JSON syntax, especially commas between list items and object properties:\n\n"
            "```json\n"
            "{\n"
            "  \"classifications\": {\n"
            "    \"[Content Piece Index String]\": [\n"
            "      [\"exact_label_name_from_valid_list_1\", 0.0],\n"
            "      [\"exact_label_name_from_valid_list_2\", 0.0]\n"
            "    ],\n"
            "    \"[Another Content Piece Index String]\": [\n"
            "      [\"exact_label_name_from_valid_list_3\", 0.0]\n"
            "    ]\n"
            "  }\n"
            "}\n"
            "```\n\n"
            "CRITICAL: Ensure every label name you output in the 'classifications' is an exact match to one of the names in the 'VALID LABEL NAMES' list provided above. Do not use descriptions or other phrases as label names. "
            "The entire response MUST be only the valid JSON object, without any surrounding text or markdown fences in the final output. Ensure all strings are double-quoted, and all lists and objects are correctly structured with necessary commas.\n"
        )

        for attempt in range(MAX_API_RETRIES + 1): # Total attempts = 1 initial + MAX_API_RETRIES
            try:
                print(f"Classification attempt {attempt + 1}/{MAX_API_RETRIES + 1} for section: \"{heading}\"")
                
                response_obj = self.model.generate_content(
                    [prompt], 
                    generation_config=GENERATION_CONFIGURATION, 
                    safety_settings=SAFETY_SETTINGS
                )
                
                # Will raise ValueError if candidate is empty/problematic (e.g. due to MAX_TOKENS, SAFETY)
                response_text = self._handle_llm_response_issues(response_obj, task_description)

                clean_response = remove_json_markdown(response_text)
                response_dict = json.loads(clean_response) # Can raise JSONDecodeError
                
                print(f"Classification successful for section: \"{heading}\" on attempt {attempt + 1}.")
                return response_dict.get("classifications", {})

            except (json.JSONDecodeError, ValueError, google_exceptions.GoogleAPIError) as e:
                # ValueError can come from _handle_llm_response_issues or direct .text access if candidate is malformed
                # GoogleAPIError for API-level issues (network, quota, server error)
                error_type = type(e).__name__
                error_message = str(e)
                print(f"Error during {task_description} on attempt {attempt + 1}/{MAX_API_RETRIES + 1}: {error_type} - {error_message}")
                
                # Log more details for certain errors if helpful
                if isinstance(e, ValueError) and "Candidate:" in error_message:
                    # The error message from _handle_llm_response_issues already includes candidate details.
                    pass 
                
                if attempt < MAX_API_RETRIES:
                    delay = RETRY_DELAY_SECONDS * (RETRY_BACKOFF_FACTOR ** attempt)
                    print(f"Retrying in {delay:.2f} seconds...")
                    time.sleep(delay)
                else:
                    final_error_message = f"{task_description} failed after {MAX_API_RETRIES + 1} attempts: {error_type} - {error_message}"
                    print(final_error_message)
                    raise RuntimeError(final_error_message) from e
        
        # This part should ideally not be reached if RuntimeError is raised on final attempt failure.
        # However, to satisfy linters or very specific control flows, returning {} is a fallback.
        print(f"Classification for section '{heading}' failed all retries and did not raise exception as expected (should not happen).")
        return {}


    def extract_target_variables(self, classified_paragraphs_data: dict) -> dict:
        """
        Extracts target variables based on classified content for each relevant tag.

        Args:
            classified_paragraphs_data (dict): Data structure from classification.
                Format: { 'tag_label': { 'heading_text': [(confidence, global_idx, content_string), ...] } }

        Returns:
            dict: Aggregated extraction results. Format: { 'variable_name': { 'value': ..., ... } }
        
        Raises:
            RuntimeError: If extraction fails for any tag_label after all retry attempts.
        """
        extraction_results = {} # This will store results across all tags
        print(f"\nStarting target variable extraction...")
        # print(f"Classified data for extraction (condensed): { {k: list(v.keys()) for k,v in classified_paragraphs_data.items()} }")


        for tag_label, headings_map in classified_paragraphs_data.items():
            task_description_for_tag = f"Extraction for tag_label '{tag_label}'"
            print(f"\nProcessing {task_description_for_tag}")

            current_target_vars_for_extraction = {} # Placeholder for your logic
            if tag_label in TARGET_VARIABLES: current_target_vars_for_extraction = {tag_label: TARGET_VARIABLES[tag_label]}
            elif tag_label in CLUSTER_TARGET_VARIABLES:
                current_target_vars_for_extraction = { var_name: TARGET_VARIABLES[var_name] for var_name in CLUSTER_TARGET_VARIABLES.get(tag_label, []) if var_name in TARGET_VARIABLES }
            if not current_target_vars_for_extraction: print(f"No target variables for tag '{tag_label}'. Skipping."); continue
            content_payload_by_heading = {}; has_content_for_this_tag = False
            for heading_text, content_tuples_list in headings_map.items():
                if heading_text not in content_payload_by_heading: content_payload_by_heading[heading_text] = {}
                for _confidence, global_idx, content_string in content_tuples_list:
                    if content_string and not content_string.isspace():
                        content_payload_by_heading[heading_text][str(global_idx)] = content_string; has_content_for_this_tag = True
            if not has_content_for_this_tag: print(f"No relevant content for tag '{tag_label}'. Skipping."); continue
            payload_for_extraction_prompt = { "target_variables_to_extract": current_target_vars_for_extraction, "headings_with_relevant_content": content_payload_by_heading }
            json_payload_for_prompt = json.dumps(payload_for_extraction_prompt, indent=2)
            prompt_start = (
    "You are an expert data extractor for systematic reviews. You are given the following JSON object. It contains:\n"
    "1. 'target_variables_to_extract': A dictionary of variables you need to extract. For each variable (the key), the value is an object containing:\n"
    "   - 'description': A detailed description of what this variable represents.\n"
    "   - 'examples': (Optional) A list of example values to guide you.\n"
    "   - 'notes_questions': (Optional) Specific notes, context, or guiding questions related to extracting this variable. You MUST consider these carefully if provided.\n"
    "2. 'headings_with_relevant_content': A dictionary where keys are section headings from a research paper. The values for each heading are dictionaries where keys are unique content piece indices (as strings) and values are the content strings (these can be text paragraphs or tables formatted as GitHub Flavored Markdown) that have been deemed relevant to the 'target_variables_to_extract' under that heading.\n\n"
    f"JSON INPUT DATA:\n```json\n{json_payload_for_prompt}\n```\n\n"
    "**CRUCIAL INSTRUCTION FOR DATA SCOPE:**\n"
    "When extracting values for the 'target_variables_to_extract', you MUST focus *exclusively* on information that describes the **primary research study** being conducted and reported in the provided content. "
    "Do NOT extract data or values that pertain to **other studies, previous work, or background literature** that are merely cited or discussed. "
    "If a target variable's specific value or detail is only found within the description of a cited study and not for the primary study's own methodology, sample, or results, then you should consider that value as 'Not Found' for the primary study.\n\n"
    "Please extract the values for the specified target variables from the provided 'headings_with_relevant_content', adhering strictly to the data scope instruction above.\n"
    "If a value cannot be found for the primary study with high confidence, set the \"value\" to 'Not Found'.\n\n"
    "For each target variable you were asked to extract (from 'target_variables_to_extract'), provide:\n"
    "- \"value\": The extracted value (string, number, boolean Y/N as appropriate, or 'Not Found').\n"
    "- \"confidence\": Your confidence in the extraction (a float from 0.0 to 1.0).\n"
    "- \"indices\": A list of a few (typically 1-5) unique content piece indices (strings, as provided in 'headings_with_relevant_content') that are **most directly relevant** to supporting your extracted 'value' and 'justification'. If the 'value' is 'Not Found' because the information is absent from the primary study, this list should ideally be empty `[]` or contain at most 1-2 indices that broadly confirm this absence.\n"
    "- \"justification\": A **very brief** explanation (preferably a single concise sentence) of how you deduced the value for the primary study, referencing specific information from the content found at the provided indices.\n\n"
    "Return your results as a single JSON object where keys are the exact variable names from the 'target_variables_to_extract' section of the input. Example format for the output JSON object structure:\n"
    "```json\n"
    "{\n"
)
            prompt_end = "\n}\n```\nYOUR ENTIRE RESPONSE MUST BE ONLY THIS VALID JSON OBJECT...\n"
            variable_output_examples = [f'  "{var_name_example}": {{"value": "[extracted value or \'Not Found\']", "confidence": 0.0, "indices": ["string_index_1"], "justification": "[brief explanation]"}}' for var_name_example in current_target_vars_for_extraction.keys()]
            prompt_middle_examples = ",\n".join(variable_output_examples)
            full_extraction_prompt = prompt_start + prompt_middle_examples + prompt_end


            for attempt in range(MAX_API_RETRIES + 1):
                try:
                    if attempt > 0: # Only print attempt number for retries
                        print(f"Extraction attempt {attempt + 1}/{MAX_API_RETRIES + 1} for tag: \"{tag_label}\"")
                    
                    response_obj = self.model.generate_content(
                        [full_extraction_prompt], 
                        generation_config=GENERATION_CONFIGURATION, 
                        safety_settings=SAFETY_SETTINGS
                    )
                    response_text = self._handle_llm_response_issues(response_obj, task_description_for_tag)
                    clean_response = remove_json_markdown(response_text)
                    response_dict = json.loads(clean_response)

                    for var_name_from_response, result_data in response_dict.items():
                        if var_name_from_response in current_target_vars_for_extraction:
                            if 'indices' in result_data and isinstance(result_data['indices'], list):
                                try:
                                    result_data['indices'] = [int(str_idx) for str_idx in result_data['indices']]
                                except ValueError:
                                    result_data['indices'] = [] 
                            else: 
                                if 'indices' not in result_data:
                                     print(f"Warning: 'indices' field missing for variable '{var_name_from_response}' under tag '{tag_label}'. Defaulting to empty list.")
                                result_data['indices'] = []
                            extraction_results[var_name_from_response] = result_data
                    
                    if attempt == 0:
                        print(f"Extraction successful for tag: \"{tag_label}\".")
                    else:
                        print(f"Extraction successful for tag: \"{tag_label}\" on attempt {attempt + 1}/{MAX_API_RETRIES + 1}.")
                    break 

                except (json.JSONDecodeError, ValueError, google_exceptions.GoogleAPIError) as e:
                    error_type = type(e).__name__
                    error_message = str(e)
                    attempt_msg_suffix = f" on attempt {attempt + 1}/{MAX_API_RETRIES + 1}" if attempt > 0 else " on initial attempt"
                    print(f"Error during {task_description_for_tag}{attempt_msg_suffix}: {error_type} - {error_message}")
                    
                    if attempt < MAX_API_RETRIES:
                        delay = RETRY_DELAY_SECONDS * (RETRY_BACKOFF_FACTOR ** attempt)
                        print(f"Retrying in {delay:.2f} seconds...")
                        time.sleep(delay)
                    else:
                        final_error_message = f"{task_description_for_tag} failed after {MAX_API_RETRIES + 1} attempts: {error_type} - {error_message}"
                        print(final_error_message)
                        raise RuntimeError(final_error_message) from e
        
        print(f"\nCompleted extraction phase. Total variables extracted: {len(extraction_results)}")
        return extraction_results

  
def remove_json_markdown(text: str) -> str:
    """Removes JSON Markdown fences from a string."""
    pattern = re.compile(r'```json\s*(.*?)\s*```', re.DOTALL)
    match = pattern.search(text)
    if match:
        return match.group(1).strip()
    return text.strip() # Return stripped original if no fences


def update_classified_data(
    indexed_content_strings: list[str],
    par_classifier_client: 'ParagraphClassifierClient', # Use forward reference if class is defined later
    classified_paragraphs_data: dict,
    current_heading: str,
    section_content_strings_for_classification: list[str],
    section_global_start_idx: int
) -> int: # Returns the count of invalid label warnings for this section
    """
    Calls the classification client for a section's content, updates the
    main classified_paragraphs_data structure, and returns the count of
    "invalid label" warnings encountered during this section's processing.

    Args:
        indexed_content_strings (list[str]): The master list of all processed content strings
                                             (paragraphs or Markdown tables) for the entire document,
                                             used for looking up content by global index.
        par_classifier_client (ParagraphClassifierClient): The client object to call the
                                                           LLM for classification.
        classified_paragraphs_data (dict): The main dictionary (accumulating results for the
                                           entire document) to update with new classifications.
                                           Format: {tag_label: {heading: [(confidence, global_idx, content_str)]}}
        current_heading (str): The text of the heading for the current section being processed.
        section_content_strings_for_classification (list[str]): The list of content strings
                                                               (paragraphs/tables) for the current section.
        section_global_start_idx (int): The global starting index in indexed_content_strings
                                        that corresponds to the first item in
                                        section_content_strings_for_classification.

    Returns:
        int: The number of "invalid label" warnings generated for this section.
    """
    invalid_label_warnings_this_section = 0

    if not section_content_strings_for_classification:
        print(f"No content strings provided for classification under heading: '{current_heading}'")
        return invalid_label_warnings_this_section

    # Get classifications from the LLM.
    # classify_section is expected to return a dictionary like:
    # {"global_idx_str": [["label1", conf1], ["label2", conf2]], ...}
    classifications = par_classifier_client.classify_section(
        current_heading,
        section_content_strings_for_classification,
        section_global_start_idx
    )

    if not classifications:
        print(f"No classifications returned from LLM for section: '{current_heading}'")
        return invalid_label_warnings_this_section

    for global_idx_str, labels_with_confidences in classifications.items():
        try:
            global_idx = int(global_idx_str)

            # Validate the global index
            if not (0 <= global_idx < len(indexed_content_strings)):
                print(f"Warning: Classified global index {global_idx} (from string '{global_idx_str}') is out of bounds "
                      f"for indexed_content_strings (length: {len(indexed_content_strings)}). Skipping this index.")
                # Incrementing warning here might be too aggressive if it's an LLM indexing error.
                # Let's assume for now this is a rare case and doesn't count towards "invalid label" warnings,
                # but rather a structural issue in the LLM response for indices.
                continue
            
            actual_content_string = indexed_content_strings[global_idx]

            if not labels_with_confidences: # If the list of labels for this index is empty
                continue

            for label_item in labels_with_confidences:
                if isinstance(label_item, list) and len(label_item) == 2:
                    label, confidence = label_item
                    try:
                        confidence_float = float(confidence)
                    except (ValueError, TypeError):
                        print(f"Warning: Could not convert confidence '{confidence}' to float for label '{label}' "
                              f"at index {global_idx}. Skipping this label-confidence pair.")
                        invalid_label_warnings_this_section += 1 # Count this as a form of invalid label data
                        continue

                    if label in PARAGRAPH_TAG_DESCRIPTIONS: # PARAGRAPH_TAG_DESCRIPTIONS from config.py
                        # Initialize nested dictionaries if they don't exist
                        if label not in classified_paragraphs_data:
                            classified_paragraphs_data[label] = {}
                        if current_heading not in classified_paragraphs_data[label]:
                            classified_paragraphs_data[label][current_heading] = []
                        
                        # Append (confidence, global_idx, actual_content_string)
                        classified_paragraphs_data[label][current_heading].append(
                            (confidence_float, global_idx, actual_content_string)
                        )
                    else:
                        # This is the specific warning the user is interested in tracking
                        print(f"Warning: Classified label '{label}' (confidence: {confidence_float}) for content index {global_idx} "
                              f"under heading '{current_heading}' is not a predefined paragraph tag. Will not be stored.")
                        invalid_label_warnings_this_section += 1
                else:
                    print(f"Warning: Malformed label-confidence pair encountered: {label_item} for index {global_idx}. Skipping.")
                    invalid_label_warnings_this_section += 1 # Count malformed pairs as warnings too

        except ValueError:
            print(f"Warning: Could not convert classified index key '{global_idx_str}' to an integer. Skipping this entry.")
            # This indicates a more structural issue with the LLM's JSON output for keys.
            # You might want a separate counter or a more severe error flag for this.
            # For now, let's count it as a type of "invalid label" related issue.
            invalid_label_warnings_this_section += 1
        except Exception as e:
            print(f"Unexpected error processing classification for index '{global_idx_str}': {e}. Skipping this entry.")
            invalid_label_warnings_this_section += 1


    if invalid_label_warnings_this_section > 0:
        print(f"Section '{current_heading}' generated {invalid_label_warnings_this_section} 'invalid label' related warnings.")
        
    return invalid_label_warnings_this_section

    
def process_document(file_path: str, par_classifier_client: 'ParagraphClassifierClient'):
    """
    Reads a Word document, converts tables to Markdown, processes content into sections
    up to (but not including) a 'Heading 2' titled 'REFERENCES', classifies them,
    and prepares data for extraction. Stops if too many invalid label warnings occur.

    Args:
        file_path (str): The path to the Word document.
        par_classifier_client (ParagraphClassifierClient): The client for classifying content.

    Returns:
        tuple: (classified_paragraphs_data, final_indexed_content_strings, final_document_content_pieces_info)
               Returns ({}, [], []) if critical error like file not found or initial processing fails.
               Can raise RuntimeError if MAX_INVALID_LABEL_WARNINGS_PER_DOC is exceeded.
    """
    print(f"Processing document: {file_path}")
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        print(f"Error opening document {file_path}: {e}")
        return {}, [], [] # Return empty structures on open failure

    # 1. Populate raw_document_content_pieces (all pieces from the doc with type, content, style)
    raw_document_content_pieces = []
    # Ensure access to python-docx objects for ._element comparison
    all_paragraphs_in_doc = list(doc.paragraphs)
    all_tables_in_doc = list(doc.tables)
    
    iter_paragraphs = iter(all_paragraphs_in_doc)
    iter_tables = iter(all_tables_in_doc)

    current_para_obj = next(iter_paragraphs, None)
    current_table_obj = next(iter_tables, None)

    # Iterate through the top-level block elements in the document body
    for block_xml_element in doc.element.body:
        if block_xml_element.tag.endswith('p'): # It's a paragraph
            if current_para_obj and block_xml_element == current_para_obj._element:
                raw_document_content_pieces.append({
                    "type": "paragraph", 
                    "content": current_para_obj.text, 
                    "style": current_para_obj.style.name
                })
                current_para_obj = next(iter_paragraphs, None)
        elif block_xml_element.tag.endswith('tbl'): # It's a table
            if current_table_obj and block_xml_element == current_table_obj._element:
                markdown_table = docx_table_to_markdown(current_table_obj)
                raw_document_content_pieces.append({
                    "type": "table_markdown", 
                    "content": markdown_table, 
                    "style": "Table" # Conceptual style name for tables
                })
                current_table_obj = next(iter_tables, None)
    
    if not raw_document_content_pieces:
        print(f"No content (paragraphs or tables) could be parsed from {file_path}.")
        return {}, [], []

    # 2. Process pieces up to "References", managing warnings
    final_indexed_content_strings = []       # List of content strings (no headings) actually processed
    final_document_content_pieces_info = [] # List of dicts for these processed content strings (for type info)

    # Assumes PARAGRAPH_TAG_DESCRIPTIONS and MAX_INVALID_LABEL_WARNINGS_PER_DOC are imported from config
    classified_paragraphs_data = {par_tag: {} for par_tag in PARAGRAPH_TAG_DESCRIPTIONS.keys()}
    current_heading_text = "Default Heading (Document Start)" 
    section_content_strings_for_classification = [] 
    current_section_actual_start_idx = 0 
    stop_processing_flag = False
    total_invalid_label_warnings_for_this_doc = 0

    for raw_piece_data in raw_document_content_pieces:
        content_string = raw_piece_data["content"]
        style_name = raw_piece_data["style"]
        
        is_heading_1 = style_name == 'Heading 1'
        is_heading_2 = style_name == 'Heading 2'
        is_any_heading = is_heading_1 or is_heading_2

        # Check for the "References" stop condition
        if is_heading_2 and content_string.strip().upper() == 'REFERENCES':
            print(f"Found '{content_string}' (Heading 2). Processing any preceding content and then stopping.")
            if section_content_strings_for_classification:
                warnings_in_section = update_classified_data(
                    final_indexed_content_strings, par_classifier_client, classified_paragraphs_data,
                    current_heading_text, section_content_strings_for_classification,
                    current_section_actual_start_idx)
                total_invalid_label_warnings_for_this_doc += warnings_in_section
                
                # Check warning threshold immediately after processing the section
                if MAX_INVALID_LABEL_WARNINGS_PER_DOC >= 0 and \
                   total_invalid_label_warnings_for_this_doc > MAX_INVALID_LABEL_WARNINGS_PER_DOC:
                    print(f"Exceeded maximum allowed invalid label warnings ({total_invalid_label_warnings_for_this_doc} > {MAX_INVALID_LABEL_WARNINGS_PER_DOC}) for document {file_path} before 'REFERENCES'.")
                    raise RuntimeError(f"Too many invalid label warnings for document {os.path.basename(file_path)}. Processing stopped.")

            stop_processing_flag = True
            break # Exit loop, do not process "References" heading or anything after

        if is_any_heading:
            # Process the previously accumulated section under the *old* heading
            if section_content_strings_for_classification:
                warnings_in_section = update_classified_data(
                    final_indexed_content_strings, par_classifier_client, classified_paragraphs_data,
                    current_heading_text, section_content_strings_for_classification,
                    current_section_actual_start_idx)
                total_invalid_label_warnings_for_this_doc += warnings_in_section

                if MAX_INVALID_LABEL_WARNINGS_PER_DOC >= 0 and \
                   total_invalid_label_warnings_for_this_doc > MAX_INVALID_LABEL_WARNINGS_PER_DOC:
                    print(f"Exceeded maximum allowed invalid label warnings ({total_invalid_label_warnings_for_this_doc} > {MAX_INVALID_LABEL_WARNINGS_PER_DOC}) for document {file_path}.")
                    raise RuntimeError(f"Too many invalid label warnings for document {os.path.basename(file_path)}. Processing stopped.")
            
            # Reset for new section, with the current piece being the new heading
            current_heading_text = content_string 
            section_content_strings_for_classification = []
            # The content for this new heading will start at the current length of final_indexed_content_strings
            current_section_actual_start_idx = len(final_indexed_content_strings)
        
        else: # It's a content piece (paragraph or table markdown)
            # If this is the first content piece for the current_heading_text
            if not section_content_strings_for_classification: 
                # Set/Confirm the starting global index for this batch of content pieces
                current_section_actual_start_idx = len(final_indexed_content_strings)
            
            section_content_strings_for_classification.append(content_string)
            
            # Add this content piece to the lists that will be returned and used for global indexing
            final_indexed_content_strings.append(content_string)
            final_document_content_pieces_info.append(raw_piece_data) 

    # After the loop, if we didn't stop early (e.g., "References" not found),
    # process the very last accumulated section.
    if not stop_processing_flag and section_content_strings_for_classification:
        warnings_in_section = update_classified_data(
            final_indexed_content_strings, par_classifier_client, classified_paragraphs_data,
            current_heading_text, section_content_strings_for_classification,
            current_section_actual_start_idx)
        total_invalid_label_warnings_for_this_doc += warnings_in_section

        if MAX_INVALID_LABEL_WARNINGS_PER_DOC >= 0 and \
           total_invalid_label_warnings_for_this_doc > MAX_INVALID_LABEL_WARNINGS_PER_DOC:
            print(f"Exceeded maximum allowed invalid label warnings ({total_invalid_label_warnings_for_this_doc} > {MAX_INVALID_LABEL_WARNINGS_PER_DOC}) for document {file_path} after processing final section.")
            raise RuntimeError(f"Too many invalid label warnings for document {os.path.basename(file_path)}. Processing stopped.")
    
    if not final_indexed_content_strings and raw_document_content_pieces: # Had raw pieces but none made it to final processing
        print(f"Content processing may have stopped very early (e.g., 'REFERENCES' at document start or all content filtered out) in {file_path}.")
    elif not raw_document_content_pieces: # Already handled this at the start, but as a final check.
        pass # Initial message about no content parsed would have been printed.

    return classified_paragraphs_data, final_indexed_content_strings, final_document_content_pieces_info


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    par_classifier_client = ParagraphClassifierClient()
    
    all_results_for_excel = []
    processing_halted_early = False
    halt_message = "" # To store the reason for halting

    try:
        print("Starting document processing. Press Control+C to interrupt and attempt to save progress.")
        files_to_process = [f for f in os.listdir(INPUT_DIR) if os.path.isfile(os.path.join(INPUT_DIR, f)) and f.lower().endswith(".docx") and not f.startswith("~")]
        
        if not files_to_process:
            print(f"No DOCX files found in the input directory: {INPUT_DIR}")
        
        for filename in files_to_process:
            file_path = os.path.join(INPUT_DIR, filename)
            print(f"\n>>> Starting processing for document: {filename}")
            
            # Functions called here (process_document, which calls client methods)
            # can raise RuntimeError after their internal retries fail.
            classified_paragraph_data, indexed_content_strings, document_content_pieces_info = \
                process_document(file_path, par_classifier_client) 
            
            if not indexed_content_strings: # Check if process_document yielded any content
                print(f"No processable content found in {filename} or processing stopped early within it. Skipping extraction for this file.")
                continue

            extracted_results = par_classifier_client.extract_target_variables(classified_paragraph_data)
            
            # Append results for the current successfully processed document
            for var_name, extraction_info in extracted_results.items():
                relevant_paragraphs_output = []
                if 'indices' in extraction_info and isinstance(extraction_info["indices"], list):
                    valid_indices = [idx for idx in extraction_info["indices"] 
                                     if isinstance(idx, int) and 0 <= idx < len(indexed_content_strings)]
                    if len(valid_indices) != len(extraction_info.get("indices", [])): # Use .get for safety
                        print(f"Warning: Some indices for variable '{var_name}' in '{filename}' were invalid or out of bounds.")
                    
                    for global_idx in valid_indices:
                        if 0 <= global_idx < len(document_content_pieces_info): # Additional check
                            content_piece_data = document_content_pieces_info[global_idx]
                            content_prefix = "[Table MD] " if content_piece_data.get("type") == "table_markdown" else "" # Use .get for safety
                            relevant_paragraphs_output.append(f"Index {global_idx}: {content_prefix}{indexed_content_strings[global_idx]}")
                        else:
                            relevant_paragraphs_output.append(f"Index {global_idx}: [Error retrieving content piece info - index out of bounds]")

                all_results_for_excel.append({
                    "filename": filename,
                    "variable": var_name,
                    "relevant_paragraphs_or_tables": "\n---\n".join(relevant_paragraphs_output),
                    "extracted_value": extraction_info.get("value", "Not Found"),
                    "confidence": extraction_info.get("confidence", 0.0),
                    "justification": extraction_info.get("justification", ""),
                    "human_verified_response": "" 
                })
            print(f"<<< Successfully processed and extracted from {filename}")

    except KeyboardInterrupt:
        print("\n\n!!! Control+C detected by user! Interrupting processing. Attempting to save progress... !!!")
        processing_halted_early = True
        halt_message = "Processing interrupted by user (Control+C)."
    except RuntimeError as e:
        # This catches errors from classify_section or extract_target_variables after retries
        print(f"\n\n!!! RUNTIME ERROR! Halting processing. Last error: {e} !!!")
        processing_halted_early = True
        halt_message = f"Critical Error: {e}"
    except Exception as e: # Catch any other truly unexpected exceptions at the main loop level
        print(f"\n\n!!! UNEXPECTED GLOBAL ERROR! Halting processing: {type(e).__name__} - {e} !!!")
        import traceback
        traceback.print_exc()
        processing_halted_early = True
        halt_message = f"Unexpected Global Error: {type(e).__name__} - {e}"
    
    finally:
        print("\n--- Finalizing run ---")
        save_file = False
        status_suffix = ""

        if processing_halted_early:
            print(f"Processing was halted: {halt_message}")
            if all_results_for_excel: # If there's any data accumulated before halt
                status_suffix = "_USER_INTERRUPTED_PARTIAL" if processing_halted_early and "Control+C" in halt_message else "_ERROR_INCOMPLETE"
                print(f"Saving partial results with suffix '{status_suffix}'.")
                save_file = True
            else: # Halted, but no results were accumulated yet
                print("No results accumulated to save.")
                # Optionally, create an empty marker file if desired, but typically not needed if no data.
        elif not all_results_for_excel: # Normal completion, but no results from any file
            print("Processing complete. No data was extracted from any document.")
        else: # Normal completion with results
            print("Processing completed successfully for all documents (or documents processed before any halt).")
            status_suffix = "_COMPLETE"
            save_file = True

        if save_file:
            df = pd.DataFrame(all_results_for_excel)
            if df.empty and not (processing_halted_early and all_results_for_excel): # Avoid saving an empty df unless it was an error with some data
                 print("DataFrame is empty and no error halt with data, not saving an empty file.")
            else:
                now = datetime.datetime.now()
                timestamp = now.strftime("%Y-%m-%d_%H-%M-%S")
                output_file = os.path.join(OUTPUT_DIR, f"extracted_data_{timestamp}{status_suffix}.xlsx")
                
                try:
                    df.to_excel(output_file, index=False)
                    print(f"Results saved to: {output_file}")
                except Exception as e_save:
                    print(f"CRITICAL: Failed to save results to Excel: {e_save}")
                    # Fallback CSV save attempt
                    csv_output_file = os.path.join(OUTPUT_DIR, f"extracted_data_{timestamp}{status_suffix}.csv")
                    try:
                        df.to_csv(csv_output_file, index=False)
                        print(f"Successfully saved results as CSV to: {csv_output_file}")
                    except Exception as e_csv_save:
                        print(f"CRITICAL: Failed to save results to CSV as fallback: {e_csv_save}")
        
        if processing_halted_early:
            print(f"Script exited due to: {halt_message}")
            sys.exit(1) # Indicate an abnormal (but handled) exit
        else:
            print("Script finished.")

if __name__ == "__main__":
    # Ensure necessary imports from config are available, e.g.,
    # from config import INPUT_DIR, OUTPUT_DIR, PARAGRAPH_TAG_DESCRIPTIONS, TARGET_VARIABLES, CLUSTER_TARGET_VARIABLES, MAX_INVALID_LABEL_WARNINGS_PER_DOC, MAX_API_RETRIES, RETRY_DELAY_SECONDS, RETRY_BACKOFF_FACTOR, GENERATION_CONFIGURATION, SAFETY_SETTINGS, PROJECT_ID, LOCATION, GEMINI_MODEL
    # This script assumes config variables are globally available after `from config import *`
    main()