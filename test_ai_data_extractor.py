import unittest
import docx
from ai_data_extractor import process_document, ParagraphClassifierClient
from config import *

@unittest.skip("temp removal")
class TestClassifySection(unittest.TestCase):
    @classmethod  
    def setUpClass(cls):  
        cls.test_doc_path = os.path.join(INPUT_DIR, "test_paper.docx") 

    def setUp(self):
        print("Starting setup...")
        self.tag_classifier_client = ParagraphClassifierClient()

    def test_classify_section(self):
        print("Starting test...")
        with open(self.test_doc_path, "rb") as docx_file:
            doc = docx.Document(docx_file)
        # Assuming the first paragraph is the heading and the second is the content
        self.assertIsNotNone(doc)
        self.assertIsNotNone(doc.paragraphs)

        heading = None
        paragraphs = None
        
        for par in doc.paragraphs:
            # Get the first heading and first paragraph under that heading
            if heading:
                paragraphs = [par.text]
                break
            if par.style.name.startswith('Heading'):
                heading = par.text
        
        # Classify the section
        classifications = self.tag_classifier_client.classify_section(heading, paragraphs, 1)  # Start index is 1
        
        # Assertions
        self.assertIsNotNone(classifications)  # Response should not be None

        paragraph_index = "1"  # Since we're passing the second paragraph (index 1)
        self.assertIn(paragraph_index, classifications)  # Paragraph index should be in the response

        tags_with_confidences = classifications[paragraph_index]
        self.assertGreater(len(tags_with_confidences), 0)  # Should have at least one tag
        for tag, confidence in tags_with_confidences:
            self.assertTrue(0 <= confidence <= 1)  # Confidence should be between 0 and 1

@unittest.skip("temp removal")
class TestProcessDocument(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.test_doc_path = os.path.join(INPUT_DIR, "test_paper.docx")
        cls.par_classifier_client = ParagraphClassifierClient()

    def test_process_document(self):
        # Process the test document
        classified_paragraphs_data = process_document(self.test_doc_path, self.par_classifier_client)
        print(f"Classified paragraphs data: {classified_paragraphs_data}")

        # Assertions
        self.assertIsNotNone(classified_paragraphs_data)                                   # Check if results are returned
        self.assertGreater(len(classified_paragraphs_data), 0)                             # Check if any variables were found
        self.assertTrue(any(classified_paragraphs_data.values()))                          # Check if at least one variable has paragraphs

        # Add assertions to validate the data structure
        for tag, heading_data in classified_paragraphs_data.items():
            assert tag in TARGET_VARIABLES or tag in CLUSTER_TARGET_VARIABLES, f"Invalid tag: {tag}"
            for heading, paragraphs in heading_data.items():
                assert isinstance(heading, str), f"Heading should be a string: {heading}"
                for confidence, index, text in paragraphs:
                    assert isinstance(confidence, float), f"Confidence should be a float: {confidence} (tag: {tag}, heading: {heading})"
                    assert 0 <= confidence <= 1, f"Confidence should be between 0 and 1: {confidence} (tag: {tag}, heading: {heading})"
                    assert isinstance(index, int), f"Index should be an integer: {index} (tag: {tag}, heading: {heading})"
                    assert isinstance(text, str), f"Text should be a string: {text} (tag: {tag}, heading: {heading})"

class TestExtractTargetVariables(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.test_doc_path = os.path.join(INPUT_DIR, "test_paper.docx")
        cls.par_classifier_client = ParagraphClassifierClient()

    def test_extract_target_variables(self):
        # Get classified paragraphs data (using the same document as TestProcessDocument)
        classified_paragraphs_data = process_document(self.test_doc_path, self.par_classifier_client)

        # Extract target variables
        extracted_results = self.par_classifier_client.extract_target_variables(classified_paragraphs_data)

        # Assertions
        self.assertIsNotNone(extracted_results)  # Check if results are returned
        self.assertGreater(len(extracted_results), 0)  # Check if any variables were extracted

        # Check the structure and content of the results (modified loop)
        for var_name, extraction_info in extracted_results.items():  # Iterate over var_name and extraction_info directly
            self.assertIn(var_name, TARGET_VARIABLES) 
            self.assertIn("value", extraction_info)
            self.assertIn("confidence", extraction_info)
            self.assertIn("indices", extraction_info)
            self.assertIn("justification", extraction_info) # Include justification check
            self.assertIsInstance(extraction_info["confidence"], (float, int))
            self.assertTrue(0 <= extraction_info["confidence"] <= 1)
            self.assertIsInstance(extraction_info["indices"], list)
            for idx in extraction_info["indices"]:
                self.assertIsInstance(idx, int)
